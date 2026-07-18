from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "SMK_CASHEW_full_stack_notes_interview_guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.1
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("SMK_CASHEW Full Stack Notes and Interview Guide")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Backend + Frontend flow, packages, middleware, debugging lessons, and interview prep")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    add_callout(
        doc,
        "How to use this guide",
        "Read Sections 1-7 for project flow and implementation understanding. Use Sections 8-10 for revision and interview preparation. The examples match the SMK_CASHEW project you built.",
    )

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "SMK_CASHEW is a MERN-style ecommerce application for selling cashew products. "
        "The backend exposes REST APIs for products, authentication, cart, checkout, and admin order management. "
        "The frontend is a Vite React app that consumes those APIs and renders customer and admin screens."
    )
    add_table(
        doc,
        ["Layer", "Main responsibility", "Important files"],
        [
            ["Backend server", "Starts Express, connects MongoDB, mounts API routes", "backend/src/server.js"],
            ["Database", "Stores users, products, carts, and orders", "MongoDB Atlas through Mongoose"],
            ["Models", "Define schemas and validation rules", "User.js, Product.js, Cart.js, Order.js"],
            ["Controllers", "Contain request handling logic", "authController, productController, cartController, checkoutController"],
            ["Routes", "Map URL + HTTP method to controllers", "authRoutes, productRoutes, cartRoutes, checkoutRoutes, adminRoutes"],
            ["Middleware", "Run reusable logic before controllers", "authMiddleware, validationMiddleware, errorMiddleware"],
            ["Frontend", "React pages, forms, navigation, and API calls", "frontend/src"],
        ],
        widths=[1.3, 2.4, 2.6],
    )

    doc.add_heading("2. End-to-End Request Flow", level=1)
    doc.add_paragraph("A typical protected request follows this path:")
    add_numbers(
        doc,
        [
            "React component triggers an action, for example clicking Add to Cart.",
            "The frontend calls the shared Axios instance from services/api.js.",
            "If the user is logged in, AuthContext sets the Authorization header as Bearer <token>.",
            "Express receives the request in server.js.",
            "The mounted route matches the path, for example /api/cart.",
            "Middleware runs first: protect verifies JWT, validation middleware checks request data, adminOnly checks role when needed.",
            "The controller runs database logic through Mongoose models.",
            "MongoDB returns data to Mongoose.",
            "Controller sends JSON response.",
            "React stores the response in state and re-renders the page.",
        ],
    )
    add_code(
        doc,
        "React page -> api.js (Axios) -> Express route -> middleware -> controller -> Mongoose model -> MongoDB -> JSON response -> React state",
    )

    doc.add_heading("3. Backend Flow", level=1)
    doc.add_heading("server.js", level=2)
    doc.add_paragraph(
        "server.js is the entry point. It loads environment variables, creates the Express app, connects MongoDB, installs global middleware, mounts route modules, and starts listening on the configured port."
    )
    add_code(
        doc,
        "app.use('/api/products', productRoutes)\napp.use('/api/auth', authRoutes)\napp.use('/api/cart', cartRoutes)\napp.use('/api/checkout', checkoutRoutes)\napp.use('/api/admin', adminRoutes)\napp.use(notFound)\napp.use(errorHandler)",
    )
    add_callout(
        doc,
        "Important ordering rule",
        "Route mounting must happen before notFound and errorHandler. If app.use('/api/admin', adminRoutes) is placed after notFound, every admin route becomes unreachable.",
    )

    doc.add_heading("Models and why mongoose.model is used", level=2)
    doc.add_paragraph(
        "A schema is only the blueprint. A model is the database interface created from that blueprint. "
        "That is why Product.js exports mongoose.model('Product', productSchema), not productSchema."
    )
    add_table(
        doc,
        ["Thing", "Meaning", "Can do database operations?"],
        [
            ["Schema", "Shape, field types, defaults, validation rules", "No"],
            ["Model", "Mongoose class connected to a MongoDB collection", "Yes: find, create, findById, update, delete"],
            ["Document", "One actual saved record", "Yes: save, validate, access fields"],
        ],
        widths=[1.3, 3.4, 1.7],
    )
    add_code(doc, "const Product = mongoose.model('Product', productSchema)\nmodule.exports = Product")

    doc.add_heading("Controllers", level=2)
    doc.add_paragraph(
        "Controllers contain the business logic for each route. They read req.params, req.query, req.body, and req.user, then call Mongoose methods and send a response."
    )
    add_table(
        doc,
        ["Controller", "Examples", "What it does"],
        [
            ["authController", "register, login, getMe", "Creates users, checks passwords, returns JWT, returns current user"],
            ["productController", "getProducts, createProduct, updateProduct", "Reads products and lets admin manage products"],
            ["cartController", "getCart, addToCart, updateCartItem, removeCartItem", "Manages one cart per user"],
            ["checkoutController", "createOrder, getMyOrders", "Creates order from cart and clears cart"],
            ["adminOrderController", "getAllOrders, updateOrderStatus", "Admin views orders and updates shipping status"],
        ],
        widths=[1.5, 2.1, 2.8],
    )

    doc.add_heading("4. Authentication and Authorization", level=1)
    doc.add_paragraph(
        "Authentication answers: who are you? Authorization answers: are you allowed to do this action?"
    )
    add_table(
        doc,
        ["Concept", "Implementation in project"],
        [
            ["JWT token", "Created during login/register with userId in payload"],
            ["protect middleware", "Reads Authorization header, verifies token, loads user from MongoDB, assigns req.user"],
            ["adminOnly middleware", "Checks req.user.role === 'admin'"],
            ["Frontend AuthContext", "Stores token in localStorage and sets Axios Authorization header"],
        ],
        widths=[2.0, 4.2],
    )
    add_code(
        doc,
        "Authorization: Bearer <token>\n\nprotect -> jwt.verify(token) -> User.findById(decoded.userId) -> req.user\nadminOnly -> req.user.role === 'admin'",
    )
    add_callout(
        doc,
        "Token-user mismatch lesson",
        "If MongoDB shows one user as admin but Postman sends a token for a different customer user, admin routes still return 403. Always verify with GET /api/auth/me using the same token.",
    )

    doc.add_heading("5. Cart and Checkout Flow", level=1)
    add_numbers(
        doc,
        [
            "Customer logs in and receives a JWT.",
            "Customer adds product to cart using POST /api/cart with productId and quantity.",
            "Cart stores product ObjectId, quantity, and priceSnapshot.",
            "GET /api/cart populates items.product so the frontend can show product details.",
            "POST /api/checkout reads the cart, verifies product availability and stock, creates an Order, decreases stock, clears the cart, and returns the created order.",
            "GET /api/checkout/my-orders returns the current user's order history.",
            "Admin uses GET /api/admin/orders and PUT /api/admin/orders/:id/status.",
        ],
    )
    add_callout(
        doc,
        "Why cart becomes empty after checkout",
        "The backend intentionally clears cart.items after creating the order. After checkout, use /api/checkout/my-orders or /api/admin/orders to see order data, not /api/cart.",
    )

    doc.add_heading("6. Frontend Flow", level=1)
    add_table(
        doc,
        ["Frontend module", "Purpose"],
        [
            ["services/api.js", "Creates one Axios instance with baseURL"],
            ["AuthContext.jsx", "Stores token/user, logs in/out, calls /auth/me"],
            ["CartContext.jsx", "Stores cart items and exposes add/update/remove/fetch cart functions"],
            ["App.jsx", "Maps URL paths to pages using React Router"],
            ["Navbar.jsx", "Shows navigation based on auth/admin state"],
            ["Products.jsx", "Fetches /products and renders ProductCard list"],
            ["AdminOrders.jsx", "Fetches /admin/orders and updates order status"],
        ],
        widths=[1.8, 4.6],
    )
    doc.add_heading("React state and effects", level=2)
    doc.add_paragraph(
        "useState stores values that affect rendering. useEffect runs side effects such as fetching API data, syncing localStorage, or reading current user information after token changes."
    )
    add_table(
        doc,
        ["Pattern", "Use when", "Example"],
        [
            ["useState(value)", "Initial state is simple", "useState('idle')"],
            ["useState(() => value)", "Initial value is read/calculated once", "useState(() => localStorage.getItem('smk_token') || '')"],
            ["setState(value)", "New value is independent of old value", "setStatus('loading')"],
            ["setState(prev => next)", "New value depends on previous value", "setItems(prev => [...prev, item])"],
            ["useEffect", "Synchronize with external systems", "Fetch API data when auth changes"],
        ],
        widths=[1.5, 2.6, 2.3],
    )
    add_callout(
        doc,
        "Effect caution",
        "Calling a setter inside useEffect is normal, but avoid creating loops where the effect depends on the same state it updates.",
    )

    doc.add_heading("7. Packages and Middleware Explained", level=1)
    add_table(
        doc,
        ["Backend package", "Why it is used"],
        [
            ["express", "Creates the HTTP server, routes, middleware pipeline, and JSON responses"],
            ["mongoose", "Defines schemas/models and talks to MongoDB"],
            ["dotenv", "Loads .env values like PORT, MONGO_URI, JWT_SECRET"],
            ["bcryptjs", "Hashes passwords and compares login password safely"],
            ["jsonwebtoken", "Creates and verifies JWT tokens"],
            ["cors", "Allows frontend dev server to call backend API"],
            ["helmet", "Adds common security-related HTTP headers"],
            ["morgan", "Logs requests during development"],
            ["express-rate-limit", "Limits repeated requests to reduce abuse"],
            ["express-validator", "Validates request body/query/params before controller logic"],
            ["nodemon", "Restarts server automatically during development"],
        ],
        widths=[1.8, 4.7],
    )
    add_table(
        doc,
        ["Frontend package", "Why it is used"],
        [
            ["vite", "Fast React development server and production build tool"],
            ["react", "Builds UI components and state-driven screens"],
            ["react-dom", "Renders React into the browser DOM"],
            ["react-router-dom", "Client-side pages such as /products, /cart, /admin/orders"],
            ["axios", "Reusable API client for backend calls"],
            ["lucide-react", "Icon library for UI controls like cart/edit/delete"],
            ["tailwindcss", "Utility classes for layout and styling"],
            ["eslint", "Finds code quality, hook, import, and React-refresh issues"],
        ],
        widths=[1.8, 4.7],
    )

    doc.add_heading("8. Debugging Lessons From This Project", level=1)
    add_table(
        doc,
        ["Symptom", "Root cause", "Fix"],
        [
            ["Mongoose URI undefined", ".env had MONG0_URI with zero instead of MONGO_URI", "Rename key to MONGO_URI"],
            ["getProductById not defined", "Function declared inside another function or export name mismatch", "Define function at top level and export same name"],
            ["JSON parse error", "Sent single quotes or JS object syntax in Postman", "Use valid JSON with double quotes"],
            ["/api/auth/me not found", "Route was POST but request was GET", "Use router.get('/me', protect, getMe)"],
            ["Cannot populate item.product", "Cart schema has items.product, not item.product", "Use populate('items.product')"],
            ["Cannot access cart before initialization", "Used lowercase cart.findOne instead of Cart.findOne", "Use model name Cart"],
            ["Admin routes 404", "Mounted admin routes after notFound middleware", "Mount routes before notFound/errorHandler"],
            ["Admin access only", "Token belonged to non-admin user", "Login with admin user and verify /auth/me"],
            ["App crash from Atlas", "Current IP not whitelisted in MongoDB Atlas", "Add IP in Atlas Network Access"],
            ["Frontend route missing", "Typo /regiter instead of /register", "Match route path and NavLink"],
        ],
        widths=[1.8, 2.4, 2.2],
    )

    doc.add_heading("9. API Reference", level=1)
    add_table(
        doc,
        ["Method", "Endpoint", "Access", "Purpose"],
        [
            ["GET", "/api/health", "Public", "Check backend running"],
            ["GET", "/api/products", "Public", "List active products"],
            ["GET", "/api/products/:id", "Public", "Get one product by Mongo _id"],
            ["POST", "/api/products", "Admin", "Create product"],
            ["PUT", "/api/products/:id", "Admin", "Update product"],
            ["DELETE", "/api/products/:id", "Admin", "Soft delete product by setting isActive false"],
            ["POST", "/api/auth/register", "Public", "Create customer user"],
            ["POST", "/api/auth/login", "Public", "Login and receive token"],
            ["GET", "/api/auth/me", "User", "Return logged-in user"],
            ["GET", "/api/cart", "User", "Return cart items"],
            ["POST", "/api/cart", "User", "Add product to cart"],
            ["PUT", "/api/cart/:productId", "User", "Update quantity"],
            ["DELETE", "/api/cart/:productId", "User", "Remove item"],
            ["POST", "/api/checkout", "User", "Create order from cart"],
            ["GET", "/api/checkout/my-orders", "User", "View own orders"],
            ["GET", "/api/admin/orders", "Admin", "View all orders"],
            ["PUT", "/api/admin/orders/:id/status", "Admin", "Update order status"],
        ],
        widths=[0.7, 2.2, 0.9, 2.7],
    )

    doc.add_heading("10. Interview Questions and Answers", level=1)
    qa = [
        ("What is the role of Express in this project?", "Express creates the backend HTTP server. It lets us define middleware, routes, request handlers, and JSON responses."),
        ("Why do we use Mongoose instead of the raw MongoDB driver?", "Mongoose gives schemas, models, validation, middleware, and convenient methods like find, create, findById, and populate."),
        ("Why export mongoose.model('Product', productSchema) instead of productSchema?", "The schema is only the blueprint. The model is the database interface that provides query and write methods."),
        ("How is _id generated if it is not in the schema?", "MongoDB automatically creates a unique ObjectId _id for every document unless explicitly disabled."),
        ("What is the difference between authentication and authorization?", "Authentication verifies who the user is. Authorization checks whether that user can perform a specific action."),
        ("How does JWT auth work in this app?", "Login creates a signed token containing userId. The frontend sends it in Authorization header. protect verifies it and loads the user."),
        ("Why should passwords be hashed?", "Plain passwords must never be stored. bcrypt stores a one-way hash, so even database exposure does not reveal raw passwords."),
        ("What does adminOnly do?", "It checks req.user.role and allows the request only when the role is exactly admin."),
        ("Why can an admin route return 403 even if one user in DB is admin?", "The token may belong to another user. The middleware checks the userId inside the token, not the user you are viewing manually."),
        ("Why use express-validator?", "It rejects invalid request data before controller/database logic, making APIs safer and easier to debug."),
        ("Why use helmet?", "Helmet sets security-related HTTP headers that reduce common web vulnerabilities."),
        ("Why use cors?", "During development, React and Express run on different origins, so CORS allows the browser to call the backend."),
        ("Why use dotenv?", "It keeps environment-specific secrets/configuration outside source code."),
        ("Why does checkout clear the cart?", "After an order is created, cart items have been converted into order items, so the active cart should be empty."),
        ("Why store priceSnapshot in cart/order items?", "Product price may change later. The cart/order should remember the price at the time of action."),
        ("What is populate in Mongoose?", "populate replaces an ObjectId reference with the referenced document data, such as turning product id into product details."),
        ("What caused the populate error item.product?", "The schema field was items.product, but the code used singular item.product."),
        ("What is React Router used for?", "It maps frontend URLs to React page components without full browser reloads."),
        ("What is Axios used for?", "Axios is the shared HTTP client used by React pages and contexts to call backend APIs."),
        ("Why use Context API for auth/cart?", "Auth and cart state are needed across many components, so Context avoids passing props through many levels."),
        ("When should you use useState callback initialization?", "When the initial value is expensive or read from storage and should run only once."),
        ("Can setState be called in useEffect?", "Yes, especially after async data fetching. Avoid infinite loops where the effect depends on the state it updates."),
        ("What does useEffect dependency array do?", "It controls when the effect re-runs based on values used by the effect."),
        ("Why did /register not render earlier?", "The route path was misspelled as /regiter while links pointed to /register."),
        ("Why did backend crash from MongoDB Atlas?", "Atlas rejected the connection because the current IP was not whitelisted."),
    ]
    for idx, (question, answer) in enumerate(qa, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{idx}. {question}")
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        doc.add_paragraph(answer)

    doc.add_heading("11. Things You Should Know Before Interviews", level=1)
    add_bullets(
        doc,
        [
            "Always explain the full request lifecycle: frontend action, API call, route, middleware, controller, model, database, response, state update.",
            "Be precise with terminology: route, controller, middleware, schema, model, document, token, header, status code.",
            "Know common status codes: 200 success, 201 created, 400 validation/client error, 401 unauthenticated, 403 forbidden, 404 not found, 500 server error.",
            "Be able to explain why environment variables and secrets should not be committed.",
            "Understand that MongoDB Atlas Network Access can break local development if your IP changes.",
            "Know how to debug systematically: read stack trace, identify file/line, inspect import/export names, reproduce with one request, fix smallest root cause.",
            "For React, know the difference between state, props, context, effects, and routing.",
            "For backend, know the difference between authentication middleware and business controllers.",
        ],
    )

    doc.add_heading("12. Quick Revision Checklist", level=1)
    add_table(
        doc,
        ["Topic", "Can you explain it?"],
        [
            ["Why mongoose.model is exported", "Schema vs model vs document"],
            ["Auth flow", "register/login -> token -> Authorization header -> protect -> req.user"],
            ["Admin flow", "protect first, adminOnly second"],
            ["Cart flow", "productId + quantity -> cart.items -> populate items.product"],
            ["Checkout flow", "cart -> order -> stock decrease -> cart cleared"],
            ["Frontend data flow", "useEffect fetch -> setState -> render"],
            ["Debugging", "Stack trace, route mount order, env variables, Atlas IP whitelist"],
        ],
        widths=[2.3, 4.1],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("SMK_CASHEW Study Guide").font.size = Pt(9)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
